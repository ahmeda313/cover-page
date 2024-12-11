from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_openai import OpenAIEmbeddings
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI
from dotenv import load_dotenv
from langchain.prompts import PromptTemplate
from langchain_community.utilities.dalle_image_generator import DallEAPIWrapper
from langchain.chains import LLMChain


from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


import requests
import os


load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")


def generate_insights(document):
    # Split the document into chunks
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
    chunks = text_splitter.split_text(document)


    # Generate embeddings for chunks and create FAISS index
    embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
#     embeddings = (
#     OpenAIEmbeddings(
#         openai_api_base=os.getenv("ST_EMBEDDINGS_BASE_URL"),
#         openai_api_key=os.getenv("PRIVATE_OPENAI_API_KEY"),
#         tiktoken_enabled=False,
#         model="mixedbread-ai/mxbai-embed-large-v1",
#     )
# )
    faiss_index = FAISS.from_texts(chunks, embeddings)
    


    # Build a RetrievalQA chain
    llm = ChatOpenAI(model="gpt-3.5-turbo", openai_api_key=openai_api_key)
    context = ''' You are an expert assistant tasked with answering questions 
            based on the most relevant context provided. Use the following pieces of 
            context to answer the question in short and as precisely as possible.
            Provide only the exact answer to my question. Do not include any extra explanation, context, or words like 'The answer is'. Keep it concise and precise.
            Example: if the question is what is the title just answer with title name instead of answering 'The title of the document is so...'
            Only the answer, please, no extra text.
            Just give answers to the point. And no negative answers.
            QUESTION : '''


    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=faiss_index.as_retriever(),

    )

    insights = []
    def ask_question(query):
        response = qa_chain.invoke(query)
        return response['result']


    insights.append(ask_question(context+"What are the key points discussed in the document? give the summary in least number of words possible"))
    insights.append(ask_question(context+"Creator of the document?, give the name also if it is an organization"))
    insights.append(ask_question(context+"Who is the document intended for?"))
    title = ask_question(context+"What is the title of the document?")
    topic = ask_question(context+"What is the most important topic in this content, give in one word?")


    image_name = image_generation(topic)

    doc_generation(title,insights[1],insights[2],insights[0])



    result = {"title":title, "image_name":image_name,"insights":insights}

    return result


def image_generation(image_description):

    # Create a prompt template to refine the image generation prompt
    print("topic", image_description)
    prompt = PromptTemplate(
        input_variables=["image_description"],
        template="Create a concise, vivid description for an image on this topic: {image_description}"
    )

    # Initialize the language model (optional, for prompt refinement)
    llm = ChatOpenAI(temperature=0.7, model="gpt-3.5-turbo")

    # Create a chain to potentially enhance the image description
    image_description_chain = LLMChain(llm=llm, prompt=prompt)

    # Generate an enhanced image description
    original_description = "give an image on this topic, it must be like a cover page {image_description}"

    # Truncate the description
    refined_description = image_description_chain.run(original_description)

    print("Refined Description:", refined_description)
    print("Description Length:", len(refined_description))

    # Use DALL-E to generate the image
    image_url = DallEAPIWrapper().run(refined_description)

    print("Generated Image URL:", image_url)

    
    response = requests.get(image_url)
    if response.status_code == 200:
        with open("./static/generated-img.png", 'wb') as f:
            f.write(response.content)
    else:
        print("Failed to download image")

    return "generated-img.png"


def doc_generation(title,from_name, to_name, summary):
    doc = Document()

    sections = doc.sections
    footer = sections[0].footer

    image_path = './static/generated-img.png'  # Replace with your image path
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(4))  # Resize as needed
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph.paragraph_format.space_before = Pt(100)

    title = doc.add_heading(title, level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.style.font.size = Pt(40)

    # key points
    paragraph = doc.add_paragraph()
    paragraph.add_run(summary)

    # Set Indentation
    paragraph.paragraph_format.left_indent = Inches(1)   # Indent from the left
    paragraph.paragraph_format.right_indent = Inches(1)  # Indent from the right

    # Align Text if Needed
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    left_para = footer.add_paragraph()
    left_para.text = "FROM\n"+from_name
    left_para.style.font.size = Pt(20)
    left_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Add bottom margin text in footer
    left_para = footer.add_paragraph()
    left_para.text = "TO\n"+to_name
    left_para.style.font.size = Pt(20)
    left_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    doc.save("./static/output.docx")